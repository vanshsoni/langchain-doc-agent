from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.chains.summarize import load_summarize_chain
from langchain_openai import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
import tempfile
import os
import logging
from dotenv import load_dotenv
import docx
import io

# Load environment variables
load_dotenv()

# Ensure OpenAI API key is set
if not os.getenv("OPENAI_API_KEY"):
    raise Exception("OPENAI_API_KEY environment variable is not set. Please check your .env file.")

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(file_bytes):
    """Extract text from TXT file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise Exception("Could not decode text file with any supported encoding")
    except Exception as e:
        raise Exception(f"Failed to extract text from TXT: {str(e)}")

def process_document(file_bytes, file_extension):
    """Process document file and create FAISS vector store"""
    temp_file = None
    temp_path = None
    docs = []
    
    try:
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            # Handle PDF files
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_file.write(file_bytes)
            temp_file.close()
            temp_path = temp_file.name
            
            logger.info("Loading PDF document...")
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
            
        elif file_extension == '.txt':
            # Handle TXT files
            text_content = extract_text_from_txt(file_bytes)
            if not text_content.strip():
                raise Exception("No text content found in TXT file")
            
            # Create a document object for consistency
            from langchain.schema import Document
            docs = [Document(page_content=text_content, metadata={"source": "text_file"})]
            
        elif file_extension in ['.doc', '.docx']:
            # Handle DOC/DOCX files
            if file_extension == '.docx':
                text_content = extract_text_from_docx(file_bytes)
            else:
                # For .doc files, we'll need to convert them first
                # This is a simplified approach - in production you might want to use a more robust solution
                raise Exception("DOC files are not supported yet. Please convert to DOCX format.")
            
            if not text_content.strip():
                raise Exception("No text content found in document")
            
            # Create a document object for consistency
            from langchain.schema import Document
            docs = [Document(page_content=text_content, metadata={"source": "document_file"})]
            
        else:
            raise Exception(f"Unsupported file type: {file_extension}")
        
        if not docs:
            raise Exception("No text content found in document")
        
        logger.info(f"Loaded document with {len(docs)} sections")
        
        # Split documents into chunks
        try:
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = splitter.split_documents(docs)
            logger.info(f"Split into {len(chunks)} chunks")
        except Exception as e:
            raise Exception(f"Failed to split text: {str(e)}")
        
        # Create embeddings and vector store
        try:
            embeddings = OpenAIEmbeddings()
            db = FAISS.from_documents(chunks, embeddings)
            logger.info("Successfully created FAISS vector store")
        except Exception as e:
            raise Exception(f"Failed to create vector store: {str(e)}")
        
        # Return both the vector store and the original documents for summarization
        return {"db": db, "docs": docs}
        
    except Exception as e:
        logger.error(f"Error in process_document: {str(e)}")
        raise
    finally:
        # Clean up temporary file
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
                logger.info("Cleaned up temporary file")
            except Exception as e:
                logger.warning(f"Failed to clean up temporary file: {str(e)}")

def summarize_document(docs):
    """Generate a summary of the document"""
    try:
        logger.info("Generating document summary...")
        
        if not docs:
            return "No content available to summarize."
        
        # Use map_reduce chain for large documents
        llm = ChatOpenAI(temperature=0)
        chain = load_summarize_chain(llm, chain_type="map_reduce")
        
        summary = chain.run(docs)
        
        logger.info("Successfully generated document summary")
        return summary
        
    except Exception as e:
        logger.error(f"Error generating summary: {str(e)}")
        return f"Could not generate summary: {str(e)}"

def create_conversation_chain(db):
    """Create a conversation chain with memory"""
    try:
        # Create memory for conversation history
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        
        # Create conversational retrieval chain
        llm = ChatOpenAI(temperature=0)
        chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=db.as_retriever(),
            memory=memory,
            return_source_documents=False,  # Don't return source documents to avoid multiple keys
            verbose=False
        )
        
        logger.info("Successfully created conversation chain with memory")
        return chain
        
    except Exception as e:
        logger.error(f"Error creating conversation chain: {str(e)}")
        raise

def ask_question(db, question, conversation_chain=None):
    """Ask a question and get answer from the vector store"""
    try:
        # Validate inputs
        if not db:
            raise Exception("No vector database provided")
        
        if not question or not question.strip():
            raise Exception("Question cannot be empty")
        
        logger.info("Retrieving relevant documents...")
        
        # Use conversation chain if available, otherwise use simple QA
        if conversation_chain:
            # Use conversational retrieval with memory
            result = conversation_chain({"question": question})
            answer = result["answer"]
            # Note: source_docs not available when return_source_documents=False
        else:
            # Fallback to simple QA without memory
            retriever = db.as_retriever()
            relevant_docs = retriever.get_relevant_documents(question)
            
            if not relevant_docs:
                return "I couldn't find any relevant information in the document to answer your question."
                
            logger.info(f"Retrieved {len(relevant_docs)} relevant documents")
            
            # Generate answer using LLM
            llm = ChatOpenAI(temperature=0)
            chain = load_qa_chain(llm, chain_type="stuff")
            answer = chain.run(input_documents=relevant_docs, question=question)
            
            if not answer or answer.strip() == "":
                return "I couldn't generate a meaningful answer to your question."
        
        logger.info("Successfully generated answer")
        return answer
        
    except Exception as e:
        logger.error(f"Error in ask_question: {str(e)}")
        raise

def get_chat_history(conversation_chain):
    """Get the current chat history"""
    if conversation_chain and hasattr(conversation_chain, 'memory'):
        return conversation_chain.memory.chat_memory.messages
    return [] 

def get_suggested_questions(db):
    """Generate suggested questions based on the document content"""
    try:
        logger.info("Generating suggested questions...")
        
        if not db:
            raise Exception("No vector database provided")
        
        # Get a few random documents to understand the content
        retriever = db.as_retriever()
        sample_docs = retriever.get_relevant_documents("", k=5)
        
        if not sample_docs:
            return ["What is this document about?", "What are the main topics?", "Can you summarize the key points?"]
        
        # Create context from sample documents
        context = "\n".join([doc.page_content for doc in sample_docs])
        
        # Generate suggested questions using LLM
        llm = ChatOpenAI(temperature=0.7)
        
        prompt = f"""Based on the following document content, generate 3 relevant and interesting questions that someone might ask about this document. 
        Make the questions specific and insightful. Return only the questions, one per line, without numbering.

        Document content:
        {context[:2000]}  # Limit context to avoid token limits

        Generate 3 questions:"""
        
        response = llm.predict(prompt)
        
        # Parse the response into individual questions
        questions = [q.strip() for q in response.split('\n') if q.strip()]
        
        # Ensure we have exactly 3 questions
        if len(questions) < 3:
            questions.extend([
                "What are the main conclusions?",
                "What are the key findings?",
                "Can you explain the important concepts?"
            ])
        
        # Take only the first 3 questions
        questions = questions[:3]
        
        logger.info("Successfully generated suggested questions")
        return questions
        
    except Exception as e:
        logger.error(f"Error generating suggested questions: {str(e)}")
        # Return fallback questions
        return [
            "What is this document about?",
            "What are the main topics discussed?",
            "Can you summarize the key points?"
        ] 